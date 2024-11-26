import os
import time
import random
import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Initialize WebDriver
driver = webdriver.Chrome()

# Set the path where images will be saved temporarily
temp_img_folder = "temp_images"
os.makedirs(temp_img_folder, exist_ok=True)

# Define the path for the Word document
doc_path = os.path.join(os.path.expanduser("~"), "Desktop", "apothecary.docx")

# Check if the document exists, and load or create it
if os.path.exists(doc_path):
    print("Document exists. Appending to it.")
    doc = Document(doc_path)  # Load the existing document
else:
    print("Document does not exist. Creating a new one.")
    doc = Document()  # Create a new document

def open_website():
    url = 'https://the-apothecary-diaries.com/'
    driver.get(url)
    driver.maximize_window()
    print(f"WebDriver successfully opened the website: {url}")
    time.sleep(5)

def gather_chapter_links():
    chapters_div = driver.find_element(By.ID, 'Chapters_List')
    chapter_links = []
    links = chapters_div.find_elements(By.TAG_NAME, 'a')
    
    for link in links:
        href = link.get_attribute('href')
        chapter_links.append(href)
    
    chapter_links.reverse()  # Reverse the list of chapters
    return chapter_links

def download_images_from_page():
    # Find all the img tags on the page
    images = driver.find_elements(By.TAG_NAME, 'img')
    img_urls = [img.get_attribute('src') for img in images]

    # Download images
    saved_images = []
    for img_url in img_urls:
        if img_url:
            # Send a GET request to download the image
            img_data = requests.get(img_url).content
            img_name = os.path.join(temp_img_folder, img_url.split("/")[-1])

            # Save the image to disk
            with open(img_name, 'wb') as img_file:
                img_file.write(img_data)
            saved_images.append(img_name)
    
    return saved_images

def add_images_to_doc(images):
    for image in images:
        doc.add_picture(image, width=Inches(1.5))  # Add the image with a size limit

def open_links_with_random_delay(chapter_links):
    for i, link in enumerate(chapter_links[78:79]):  # Process chapters 31 to 60
        driver.get(link)
        print(f"Opened: {link}")
        
        # Download images from the page
        images = download_images_from_page()
        
        # Count the number of images
        image_count = len(images)

        print(f"Number of images: {image_count}")
        
        # Add the images to the Word document
        add_images_to_doc(images)
        
        print(f"Waiting for 8 seconds...")
        time.sleep(8)

# Main execution
open_website()

# Gather and reverse chapter links from the website
chapter_links = gather_chapter_links()

# Open links, download images, and append to the document
open_links_with_random_delay(chapter_links)

# Save the Word document
doc.save(doc_path)
print(f"Word document saved at: {doc_path}")

# Cleanup: Remove the temporary images folder and its content
import shutil
shutil.rmtree(temp_img_folder)

driver.quit()
